from __future__ import annotations
import io
import re
from dataclasses import dataclass
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
from agents.models import Outline
from export.template_reader import TemplateStyles
from images.generator import GeneratedImage

TARGET_WIDTH_PX = int(4.0 * 96)  # 4 inches at 96 DPI

# Tokeniser for inline markdown: bold (**text**), italic (*text*), plain text
_INLINE_TOKEN = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)", re.DOTALL)


def _add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _plain_para(doc: Document, text: str, font_size: float = 11,
                bold: bool = False, italic: bool = False,
                align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Add a paragraph with explicit Garamond formatting — no style names."""
    para = doc.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.font.name = "Garamond"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic


def _markdown_para(doc: Document, text: str, font_size: float = 11,
                   align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Add a paragraph, rendering **bold** and *italic* inline markdown as runs."""
    para = doc.add_paragraph()
    para.alignment = align
    for m in _INLINE_TOKEN.finditer(text):
        bold_text, italic_text, plain_text = m.group(1), m.group(2), m.group(3)
        if bold_text:
            run = para.add_run(bold_text)
            run.font.name = "Garamond"
            run.font.size = Pt(font_size)
            run.font.bold = True
        elif italic_text:
            run = para.add_run(italic_text)
            run.font.name = "Garamond"
            run.font.size = Pt(font_size)
            run.font.italic = True
        elif plain_text:
            run = para.add_run(plain_text)
            run.font.name = "Garamond"
            run.font.size = Pt(font_size)


@dataclass
class DocxBuilder:
    styles: TemplateStyles

    def build(
        self,
        outline: Outline,
        chapter_texts: dict[int, str],
        chapter_images: dict[int, list[GeneratedImage]],
        about_author: str,
    ) -> bytes:
        doc = Document()

        for section in doc.sections:
            section.page_width = Inches(self.styles.page_width_inches)
            section.page_height = Inches(self.styles.page_height_inches)
            section.top_margin = Inches(self.styles.margin_top_inches)
            section.bottom_margin = Inches(self.styles.margin_bottom_inches)
            section.left_margin = Inches(self.styles.margin_left_inches)
            section.right_margin = Inches(self.styles.margin_right_inches)

        self._add_title_page(doc, outline)
        _add_page_break(doc)
        self._add_copyright(doc, outline)
        _add_page_break(doc)
        self._add_front_matter_section(doc, "DEDICATION", "For all curious minds.")
        _add_page_break(doc)
        self._add_front_matter_section(doc, "ACKNOWLEDGMENTS", "")
        _add_page_break(doc)
        self._add_contents(doc, outline)
        _add_page_break(doc)

        for chapter in outline.chapters:
            text = chapter_texts.get(chapter.number, "")
            images = chapter_images.get(chapter.number, [])
            self._add_chapter(doc, chapter, text, images)
            _add_page_break(doc)

        self._add_about_author(doc, about_author)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_title_page(self, doc: Document, outline: Outline) -> None:
        _plain_para(doc, outline.book_title, font_size=36,
                    align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        _plain_para(doc, outline.author, font_size=18,
                    align=WD_ALIGN_PARAGRAPH.CENTER)

    def _add_copyright(self, doc: Document, outline: Outline) -> None:
        year = datetime.now().year
        lines = [
            f"Copyright © {year} {outline.author}",
            "All rights reserved.",
            "",
            "No part of this publication may be reproduced, distributed, or transmitted "
            "in any form or by any means without the prior written permission of the author.",
        ]
        for line in lines:
            _plain_para(doc, line, font_size=10)

    def _add_front_matter_section(self, doc: Document, heading: str, body: str) -> None:
        _plain_para(doc, heading, font_size=14, bold=True)
        doc.add_paragraph()
        if body:
            _plain_para(doc, body, font_size=11)

    def _add_contents(self, doc: Document, outline: Outline) -> None:
        _plain_para(doc, "CONTENTS", font_size=14, bold=True)
        doc.add_paragraph()
        for ch in outline.chapters:
            _plain_para(doc, f"{ch.number}. {ch.title}", font_size=11)

    def _add_chapter(
        self,
        doc: Document,
        chapter,
        text: str,
        images: list[GeneratedImage],
    ) -> None:
        _plain_para(doc, f"{chapter.number}  {chapter.title.upper()}",
                    font_size=14, bold=True)
        doc.add_paragraph()

        # Strip leading title that Haiku typically repeats at the top of the chapter.
        # Matches: "# Title", "**Title**", "Chapter N: Title", plain title (any case).
        title_pattern = re.compile(
            r"^\s*"
            r"(?:\*{1,2})?"
            r"(?:#{1,6}\s*)?"
            r"(?:chapter\s+\d+\s*[:\-–]\s*)?"
            r"(?:\*{1,2})?"
            r"(?:\d+\s+)?"
            + re.escape(chapter.title)
            + r"(?:\*{1,2})?"
            + r"\s*\n?",
            re.IGNORECASE,
        )
        cleaned_text = title_pattern.sub("", text, count=1).lstrip()

        image_map: dict[str, GeneratedImage] = {
            img.marker.full_match: img for img in images
        }
        image_pattern = re.compile(r"(\[IMAGE:[^\]]+\])")
        parts = image_pattern.split(cleaned_text)

        seen: set[str] = set()
        is_first_para = True

        for part in parts:
            if image_pattern.match(part):
                img = image_map.get(part)
                if img:
                    self._embed_image(doc, img)
                else:
                    _plain_para(doc, part, font_size=10, italic=True)
            else:
                for para_text in self._split_paragraphs(part):
                    key = re.sub(r"\s+", " ", para_text.lower())
                    if key in seen:
                        continue
                    seen.add(key)

                    # Section headings (stripped of ##) rendered as bold body text
                    is_heading = bool(re.match(r"^#{1,6}\s", para_text))
                    clean = re.sub(r"^#{1,6}\s+", "", para_text)

                    if is_heading:
                        _plain_para(doc, clean, font_size=11, bold=True)
                    elif is_first_para:
                        _markdown_para(doc, clean, font_size=11)
                        is_first_para = False
                    else:
                        _markdown_para(doc, clean, font_size=11)

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        """Split on blank lines; join soft-wrapped single newlines into one paragraph."""
        blocks = re.split(r"\n{2,}", text)
        return [b.replace("\n", " ").strip() for b in blocks if b.strip()]

    def _embed_image(self, doc: Document, img: GeneratedImage) -> None:
        buf = io.BytesIO(img.image_bytes)
        pil_img = PILImage.open(buf)
        w, h = pil_img.size
        if w > TARGET_WIDTH_PX:
            new_h = int(h * TARGET_WIDTH_PX / w)
            pil_img = pil_img.resize((TARGET_WIDTH_PX, new_h), PILImage.LANCZOS)
        buf_out = io.BytesIO()
        pil_img.save(buf_out, format="PNG")
        buf_out.seek(0)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(buf_out, width=Inches(4.0))

        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(img.marker.description)
        caption_run.italic = True
        caption_run.font.name = "Garamond"
        caption_run.font.size = Pt(9)

    def _add_about_author(self, doc: Document, bio: str) -> None:
        _plain_para(doc, "ABOUT THE AUTHOR", font_size=14, bold=True)
        doc.add_paragraph()
        _plain_para(doc, bio, font_size=11)
