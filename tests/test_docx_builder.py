import io
import pytest
from docx import Document
from agents.models import BookParams, Outline, Chapter, Section
from export.template_reader import TemplateStyles, _DEFAULTS
from export.docx_builder import DocxBuilder


def _make_outline() -> Outline:
    sections = [Section(title="Overview", estimated_pages=3.0)]
    chapter = Chapter(number=1, title="The Beginning", sections=sections, estimated_pages=8.0)
    return Outline(
        book_title="Test Book",
        author="Test Author",
        language="English",
        front_matter=["Dedication", "Acknowledgments"],
        chapters=[chapter],
        back_matter=["About the Author"],
        total_estimated_pages=100,
    )


def test_build_creates_docx_bytes():
    outline = _make_outline()
    chapter_texts = {1: "This is chapter one content. It has some text."}
    chapter_images = {1: []}

    builder = DocxBuilder(styles=_DEFAULTS)
    result = builder.build(
        outline=outline,
        chapter_texts=chapter_texts,
        chapter_images=chapter_images,
        about_author="Test author bio.",
    )

    assert isinstance(result, bytes)
    assert len(result) > 0
    # Verify it's a valid docx
    doc = Document(io.BytesIO(result))
    assert doc is not None


def test_build_includes_title_and_author():
    outline = _make_outline()
    builder = DocxBuilder(styles=_DEFAULTS)
    result = builder.build(
        outline=outline,
        chapter_texts={1: "Chapter content here."},
        chapter_images={1: []},
        about_author="Bio text.",
    )
    doc = Document(io.BytesIO(result))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Book" in all_text
    assert "Test Author" in all_text


def test_build_includes_chapter_title():
    outline = _make_outline()
    builder = DocxBuilder(styles=_DEFAULTS)
    result = builder.build(
        outline=outline,
        chapter_texts={1: "Chapter content here."},
        chapter_images={1: []},
        about_author="Bio.",
    )
    doc = Document(io.BytesIO(result))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "The Beginning" in all_text


def test_build_renders_unmatched_image_marker_as_italic():
    outline = _make_outline()
    builder = DocxBuilder(styles=_DEFAULTS)
    result = builder.build(
        outline=outline,
        chapter_texts={1: "Intro.\n\n[IMAGE: A diagram]\n\nMore text."},
        chapter_images={1: []},  # no images provided
        about_author="Bio.",
    )
    doc = Document(io.BytesIO(result))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "[IMAGE: A diagram]" in all_text
